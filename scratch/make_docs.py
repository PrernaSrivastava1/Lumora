import docx

def create_docx():
    doc = docx.Document()
    doc.add_heading('Deep Learning and NLP', level=1)
    doc.add_paragraph('Natural Language Processing (NLP) is a subfield of artificial intelligence. It focuses on the interaction between computers and human language. Modern NLP applications rely on neural network architectures like Transformers to understand syntax, semantics, and context. By converting words and sentences into high-dimensional vector embeddings, models can perform tasks such as machine translation, sentiment analysis, and question answering. Algorithms like KD-Trees and HNSW allow rapid retrieval of these embeddings from vector databases, forming the backbone of Retrieval-Augmented Generation (RAG) systems.')
    doc.save('test.docx')
    print("Created test.docx")

if __name__ == "__main__":
    create_docx()
